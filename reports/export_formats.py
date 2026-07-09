"""
Feature 11: Export Formats (DOCX and plain text, alongside the existing PDF).

Deliberately built as thin siblings of pdf_report.py rather than a shared
abstract renderer: the three formats have different enough native structures
(DOCX has real paragraph/heading objects, plain text is just strings, PDF
uses flowables) that a forced common abstraction would add complexity
without saving much code. All three consume the exact same `report_data`
dict shape, so content stays consistent across formats - that's the
consistency Feature 11 asks for, not identical rendering code.

Install:
    pip install python-docx --break-system-packages
"""
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pdf_report import _apply_privacy_redaction


def build_docx(report_data: dict, output_path: str):
    if report_data.get("privacy_mode"):
        report_data = _apply_privacy_redaction(dict(report_data))

    doc = Document()

    title = doc.add_heading("Cybercrime Incident Report", level=0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Generated on {datetime.now().strftime('%d %B %Y, %I:%M %p')} — for informational "
        f"and complaint-filing purposes only, not a legal document by itself."
    ).italic = True

    if report_data.get("is_urgent"):
        alert = doc.add_paragraph()
        run = alert.add_run(f"⚠ URGENT: {report_data.get('emergency_message', '')}")
        run.bold = True
        run.font.color.rgb = RGBColor(0x7A, 0x12, 0x12)

    doc.add_heading("Complaint Summary", level=1)
    summary_pairs = [
        ("Detected Crime Type", report_data.get("crime_type", "").replace("_", " ").title()),
        ("AI Confidence Score", f"{report_data.get('confidence', 0):.0%}"),
        ("Date & Time of Incident", report_data.get("incident_datetime", "")),
        ("Platform Involved", report_data.get("platform", "")),
        ("Financial Loss", report_data.get("financial_loss", "")),
    ]
    table = doc.add_table(rows=0, cols=2)
    for label, value in summary_pairs:
        if value:
            row = table.add_row().cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].bold = True
            row[1].text = str(value)

    if report_data.get("regime_note"):
        note = doc.add_paragraph()
        note.add_run(report_data["regime_note"]).italic = True

    doc.add_heading("Incident Description", level=2)
    doc.add_paragraph(report_data.get("user_text", ""))

    victim = report_data.get("victim_details") or {}
    if any(victim.values()):
        doc.add_heading("Victim Details", level=1)
        for label, key in [("Name", "name"), ("Phone", "phone"), ("Email", "email"), ("Address", "address")]:
            if victim.get(key):
                doc.add_paragraph(f"{label}: {victim[key]}")

    suspect = report_data.get("suspect_details") or {}
    if any(suspect.values()):
        doc.add_heading("Suspect Details", level=1)
        for label, key in [("Name", "name"), ("Phone", "phone"), ("Email", "email"),
                            ("UPI ID", "upi_id"), ("Bank Account", "bank_account")]:
            if suspect.get(key):
                doc.add_paragraph(f"{label}: {suspect[key]}")

    if report_data.get("device_info") or report_data.get("ip_address"):
        doc.add_heading("Device Information", level=1)
        if report_data.get("device_info"):
            doc.add_paragraph(f"Device Info: {report_data['device_info']}")
        if report_data.get("ip_address"):
            doc.add_paragraph(f"IP Address: {report_data['ip_address']}")

    entities = report_data.get("extracted_entities") or {}
    if any(entities.values()):
        doc.add_heading("Extracted Information", level=1)
        for label, key in [
            ("Phone Numbers", "phone_numbers"), ("Email Addresses", "emails"),
            ("URLs", "urls"), ("Bank Details", "possible_bank_accounts"),
            ("UPI IDs", "upi_ids"), ("Transaction IDs", "possible_transaction_ids"),
        ]:
            values = entities.get(key)
            if values:
                doc.add_paragraph(f"{label}: {', '.join(values)}")

    evidence = report_data.get("evidence_list") or []
    if evidence:
        doc.add_heading("Uploaded Evidence", level=1)
        ev_table = doc.add_table(rows=1, cols=4)
        hdr = ev_table.rows[0].cells
        for i, h in enumerate(["File Name", "Type", "Uploaded", "SHA-256"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for e in evidence:
            row = ev_table.add_row().cells
            row[0].text = e.get("filename", "")
            row[1].text = e.get("file_type", "")
            row[2].text = e.get("upload_time", "")
            row[3].text = (e.get("sha256", "")[:16] + "…") if e.get("sha256") else ""

    timeline = report_data.get("timeline") or []
    if timeline:
        doc.add_heading("AI-Generated Timeline", level=1)
        for t in timeline:
            p = doc.add_paragraph()
            p.add_run(f"{t.get('time', '')}").bold = True
            p.add_run(f" — {t.get('event', '')}")

    applicable_law = report_data.get("applicable_law") or []
    if applicable_law:
        doc.add_heading("Relevant Legal Sections", level=1)
        for law in applicable_law:
            p = doc.add_paragraph()
            p.add_run(f"[{law.get('chunk_id', '')}] ").bold = True
            p.add_run(law.get("plain_language_summary", ""))

    actions = report_data.get("immediate_actions") or []
    if actions:
        doc.add_heading("Recommended Next Steps", level=1)
        for a in actions:
            doc.add_paragraph(a, style="List Bullet")

    safety = report_data.get("safety_recommendations") or []
    if safety:
        doc.add_heading("Safety Recommendations", level=1)
        for s in safety:
            doc.add_paragraph(s, style="List Bullet")

    draft = report_data.get("draft_complaint") or ""
    if draft:
        doc.add_heading("Draft Complaint (for reference)", level=1)
        doc.add_paragraph(draft)

    if report_data.get("uncovered_aspects"):
        doc.add_heading("Additional Notes", level=1)
        doc.add_paragraph(report_data["uncovered_aspects"])

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "This report was generated with AI assistance and is intended to help you organize "
        "information for filing a complaint. It does not constitute legal advice."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)

    doc.save(output_path)
    return output_path


def build_plain_text(report_data: dict) -> str:
    if report_data.get("privacy_mode"):
        report_data = _apply_privacy_redaction(dict(report_data))

    lines = []
    lines.append("CYBERCRIME INCIDENT REPORT")
    lines.append(f"Generated on {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    lines.append("(Informational and complaint-filing purposes only, not a legal document by itself.)")
    lines.append("")

    if report_data.get("is_urgent"):
        lines.append(f"*** URGENT: {report_data.get('emergency_message', '')} ***")
        lines.append("")

    lines.append("-- COMPLAINT SUMMARY --")
    for label, value in [
        ("Detected Crime Type", report_data.get("crime_type", "").replace("_", " ").title()),
        ("AI Confidence Score", f"{report_data.get('confidence', 0):.0%}"),
        ("Date & Time of Incident", report_data.get("incident_datetime", "")),
        ("Platform Involved", report_data.get("platform", "")),
        ("Financial Loss", report_data.get("financial_loss", "")),
    ]:
        if value:
            lines.append(f"{label}: {value}")
    if report_data.get("regime_note"):
        lines.append(report_data["regime_note"])
    lines.append("")
    lines.append("Incident Description:")
    lines.append(report_data.get("user_text", ""))
    lines.append("")

    victim = report_data.get("victim_details") or {}
    if any(victim.values()):
        lines.append("-- VICTIM DETAILS --")
        for label, key in [("Name", "name"), ("Phone", "phone"), ("Email", "email"), ("Address", "address")]:
            if victim.get(key):
                lines.append(f"{label}: {victim[key]}")
        lines.append("")

    suspect = report_data.get("suspect_details") or {}
    if any(suspect.values()):
        lines.append("-- SUSPECT DETAILS --")
        for label, key in [("Name", "name"), ("Phone", "phone"), ("Email", "email"),
                            ("UPI ID", "upi_id"), ("Bank Account", "bank_account")]:
            if suspect.get(key):
                lines.append(f"{label}: {suspect[key]}")
        lines.append("")

    if report_data.get("device_info") or report_data.get("ip_address"):
        lines.append("-- DEVICE INFORMATION --")
        if report_data.get("device_info"):
            lines.append(f"Device Info: {report_data['device_info']}")
        if report_data.get("ip_address"):
            lines.append(f"IP Address: {report_data['ip_address']}")
        lines.append("")

    entities = report_data.get("extracted_entities") or {}
    if any(entities.values()):
        lines.append("-- EXTRACTED INFORMATION --")
        for label, key in [
            ("Phone Numbers", "phone_numbers"), ("Email Addresses", "emails"),
            ("URLs", "urls"), ("Bank Details", "possible_bank_accounts"),
            ("UPI IDs", "upi_ids"), ("Transaction IDs", "possible_transaction_ids"),
        ]:
            values = entities.get(key)
            if values:
                lines.append(f"{label}: {', '.join(values)}")
        lines.append("")

    evidence = report_data.get("evidence_list") or []
    if evidence:
        lines.append("-- UPLOADED EVIDENCE --")
        for e in evidence:
            lines.append(f"- {e.get('filename', '')} ({e.get('file_type', '')}, uploaded {e.get('upload_time', '')}, SHA-256: {e.get('sha256', '')})")
        lines.append("")

    timeline = report_data.get("timeline") or []
    if timeline:
        lines.append("-- AI-GENERATED TIMELINE --")
        for t in timeline:
            lines.append(f"{t.get('time', '')} — {t.get('event', '')}")
        lines.append("")

    applicable_law = report_data.get("applicable_law") or []
    if applicable_law:
        lines.append("-- RELEVANT LEGAL SECTIONS --")
        for law in applicable_law:
            lines.append(f"[{law.get('chunk_id', '')}] {law.get('plain_language_summary', '')}")
        lines.append("")

    actions = report_data.get("immediate_actions") or []
    if actions:
        lines.append("-- RECOMMENDED NEXT STEPS --")
        for a in actions:
            lines.append(f"- {a}")
        lines.append("")

    safety = report_data.get("safety_recommendations") or []
    if safety:
        lines.append("-- SAFETY RECOMMENDATIONS --")
        for s in safety:
            lines.append(f"- {s}")
        lines.append("")

    draft = report_data.get("draft_complaint") or ""
    if draft:
        lines.append("-- DRAFT COMPLAINT (FOR REFERENCE) --")
        lines.append(draft)
        lines.append("")

    if report_data.get("uncovered_aspects"):
        lines.append("-- ADDITIONAL NOTES --")
        lines.append(report_data["uncovered_aspects"])
        lines.append("")

    lines.append(
        "This report was generated with AI assistance and is intended to help you organize "
        "information for filing a complaint. It does not constitute legal advice."
    )
    return "\n".join(lines)


if __name__ == "__main__":
    sample = {
        "crime_type": "upi_fraud", "confidence": 0.91,
        "user_text": "Someone took money from my bank account using a fake UPI link.",
        "platform": "PhonePe", "financial_loss": "Rs 15,000",
        "is_urgent": True, "emergency_message": "Call your bank and 1930 immediately.",
        "applicable_law": [{"chunk_id": "it_act_66d", "plain_language_summary": "Covers cheating by personation."}],
        "immediate_actions": ["Call your bank's fraud helpline."],
        "safety_recommendations": ["Change your UPI PIN."],
        "draft_complaint": "I am writing to report an unauthorized UPI transaction...",
        "timeline": [{"time": "09:20 AM", "event": "Money was debited."}],
    }
    build_docx(sample, "/tmp/sample_report.docx")
    print("DOCX built at /tmp/sample_report.docx")
    text = build_plain_text(sample)
    with open("/tmp/sample_report.txt", "w") as f:
        f.write(text)
    print("Plain text built at /tmp/sample_report.txt")
    print()
    print(text[:300], "...")
